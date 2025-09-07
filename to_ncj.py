#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX to Normalized Content JSON (NCJ) converter with advanced figure grouping.

Direct python-docx based implementation for accurate figure grouping and 
title/credit attribution.

Usage:
  python to_ncj.py "input.docx" --out content.json [options]
"""
import sys, json, re, os, hashlib, zipfile, shutil
from typing import List, Dict, Any, Tuple, Optional, NamedTuple
from dataclasses import dataclass
from docx import Document
from docx.shared import Emu
import opencc

@dataclass
class FigureCandidate:
    """Represents a potential figure found in the document"""
    para_idx: int
    run_idx: int  
    width_emu: Optional[int] = None
    height_emu: Optional[int] = None
    r_id: Optional[str] = None  # rId from OOXML relationships
    media_path: Optional[str] = None  # path in ZIP: word/media/imageN.ext
    asset_id: Optional[str] = None  # SHA256-based asset ID
    filename: Optional[str] = None  # extracted file path

@dataclass
class GroupCandidate:
    """Represents a potential group of figures"""
    figures: List[FigureCandidate]
    layout: str  # 'row' or 'column'
    title: Optional[str] = None
    credit: Optional[str] = None
    reason: Optional[str] = None  # Debug info

# --------- Configuration and Helpers ----------
class Config:
    def __init__(self):
        self.max_title_len = 60
        self.max_gap_paras = 1
        self.page_width_ratio = 0.95
        self.debug = False
        self.assets_dir = "assets/media"
        self.traditional_to_simplified = False  # Enable traditional to simplified conversion
        self.use_explicit_markers = True  # Use explicit title/source markers for better accuracy

def safe_traditional_to_simplified(text: str, converter) -> str:
    """安全的繁简转换，保护特定模式"""
    if not text or not text.strip():
        return text
    
    # 保护模式：先替换为占位符
    protected_replacements = []
    temp_text = text
    
    # 保护拉丁字符（包括单词、缩写、网站等）
    latin_pattern = r'\b[a-zA-Z][a-zA-Z0-9\-\.]*[a-zA-Z0-9]\b|\b[a-zA-Z]\b'
    latin_matches = re.findall(latin_pattern, temp_text)
    for i, match in enumerate(latin_matches):
        placeholder = f'___LATIN_{i:03d}___'
        protected_replacements.append((placeholder, match))
        temp_text = temp_text.replace(match, placeholder, 1)
    
    # 保护数字和百分比
    number_pattern = r'\d+\.?\d*%?'
    number_matches = re.findall(number_pattern, temp_text)
    for i, match in enumerate(number_matches):
        placeholder = f'___NUMBER_{i:03d}___'
        protected_replacements.append((placeholder, match))
        temp_text = temp_text.replace(match, placeholder, 1)
    
    # 保护URL和邮箱
    url_pattern = r'https?://[^\s]+|\b[\w\.-]+@[\w\.-]+\.[a-zA-Z]{2,}\b'
    url_matches = re.findall(url_pattern, temp_text)
    for i, match in enumerate(url_matches):
        placeholder = f'___URL_{i:03d}___'
        protected_replacements.append((placeholder, match))
        temp_text = temp_text.replace(match, placeholder, 1)
    
    # 执行繁简转换
    try:
        simplified_text = converter.convert(temp_text)
    except Exception:
        # 转换失败时返回原文
        simplified_text = temp_text
    
    # 恢复保护的内容
    for placeholder, original in reversed(protected_replacements):
        simplified_text = simplified_text.replace(placeholder, original, 1)
    
    return simplified_text

def sha256_of_file(path: str) -> str:
    h = hashlib.sha256()
    with open(path, 'rb') as f:
        for chunk in iter(lambda: f.read(65536), b''):
            h.update(chunk)
    return h.hexdigest()

def normalize_credit(text: str) -> str:
    if not text: return ""
    # Remove 来源:/來源:/Source: (both half-width and full-width colons)
    text = re.sub(r'^\s*(来源|來源|Source)\s*[：:]\s*', '', text, flags=re.I)
    # Remove leading/trailing spaces and ending punctuation
    text = text.strip().rstrip('.,;，。；')
    # Add unified "Source: " prefix
    if text:
        text = f"Source: {text}"
    return text

def is_short_title(text: str, max_len: int) -> bool:
    """Check if text could be a figure title"""
    if not text or len(text.strip()) == 0:
        return False
    if len(text.strip()) > max_len:
        return False
    # Exclude obvious credit lines
    if re.match(r'^\s*(来源|來源|Source)\s*[：:]', text, re.I):
        return False
    return True

def is_credit_line(text: str) -> bool:
    """Check if text is a credit/source line (legacy method)"""
    return bool(re.match(r'^\s*(来源|來源|Source)\s*[：:]', text, re.I))

def is_potential_credit_line(text: str, max_len: int = 60) -> bool:
    """更严格但增强的来源检测 - 仅使用传统严格模式"""
    if not text or len(text.strip()) == 0:
        return False
    if len(text.strip()) > max_len:  # 来源信息通常较短
        return False
    
    # 传统严格模式：以"来源:/Source:"开头
    if re.match(r'^\s*(来源|來源|Source)\s*[：:]', text, re.I):
        return True
    
    return False

def is_title_paragraph(text: str) -> bool:
    """检测是否为明确标记的标题段落"""
    if not text:
        return False
    text_cleaned = text.strip().lower()
    return (text_cleaned.startswith('标题') or 
            text_cleaned.startswith('title'))

def is_source_paragraph(text: str) -> bool:
    """检测是否为明确标记的来源段落"""  
    if not text:
        return False
    text_cleaned = text.strip().lower()
    return (text_cleaned.startswith('来源') or 
            text_cleaned.startswith('來源') or
            text_cleaned.startswith('source'))

def extract_title_content(text: str) -> str:
    """提取标题内容（去除前缀标记）"""
    if not text:
        return ""
    text = text.strip()
    
    # 定义所有可能的标题前缀（按长度排序，长的优先）
    prefixes = [
        '标题：', '标题:', 'Title：', 'Title:', 
        '标题 ', 'Title ', '标题', 'Title'
    ]
    
    for prefix in prefixes:
        if text.lower().startswith(prefix.lower()):
            return text[len(prefix):].strip()
    
    return text

def extract_source_content(text: str) -> str:
    """提取来源内容（去除前缀标记，统一添加Source前缀）"""
    if not text:
        return ""
    text = text.strip()
    
    # 定义所有可能的来源前缀（按长度排序，长的优先）
    prefixes = [
        '来源：', '来源:', '來源：', '來源:', 'Source：', 'Source:',
        '来源 ', '來源 ', 'Source ', '来源', '來源', 'Source'
    ]
    
    for prefix in prefixes:
        if text.lower().startswith(prefix.lower()):
            content = text[len(prefix):].strip()
            # Add unified "Source: " prefix
            return f"Source: {content}" if content else ""
    
    return text

DOC_TITLE_RE = re.compile(r'^\s*(\d{6})\s*-\s*(.+)')
def parse_date_from_yyMMdd(yyMMdd: str) -> Optional[str]:
    try:
        yy = int(yyMMdd[0:2]); mm = int(yyMMdd[2:4]); dd = int(yyMMdd[4:6])
        yyyy = 2000 + yy
        return f"{yyyy:04d}-{mm:02d}-{dd:02d}"
    except Exception:
        return None

# --------- Phase 1: Extract Figure Candidates ----------
def extract_figures_from_docx(doc: Document, docx_path: str, converter=None) -> Tuple[List[FigureCandidate], List[str], int]:
    """Extract figure candidates and all text content (paragraphs + tables) with real rId mapping"""
    figures = []
    para_texts = []
    page_width_emu = doc.sections[0].page_width.emu if doc.sections else 7559675
    
    # Build rId to media path mapping from document relations
    r_id_to_media = {}
    for r_id, rel in doc.part._rels.items():
        if rel.reltype.endswith('/image'):  # Image relationship
            r_id_to_media[r_id] = rel.target_ref  # e.g., 'media/image1.png'
    
    # Extract all content in document order (paragraphs + tables)
    body = doc._element.body
    content_idx = 0
    paragraph_counter = 0
    
    for element in body:
        if element.tag.endswith('}p'):  # Paragraph
            # Find corresponding paragraph object
            if paragraph_counter < len(doc.paragraphs):
                para = doc.paragraphs[paragraph_counter]
                para_text = para.text.strip()
                if converter:
                    para_text = safe_traditional_to_simplified(para_text, converter)
                para_texts.append(para_text)
                
                # Find drawings in this paragraph
                drawings = para._element.xpath('.//w:drawing')
                for run_idx, drawing in enumerate(drawings):
                    # Extract dimensions if available
                    extents = drawing.xpath('.//wp:extent')
                    width_emu = height_emu = None
                    if extents:
                        extent = extents[0]
                        width_emu = int(extent.get('cx', 0))
                        height_emu = int(extent.get('cy', 0))
                    
                    # Extract rId from drawing XML
                    r_id = None
                    media_path = None
                    try:
                        xml_str = drawing.xml
                        r_embed_matches = re.findall(r'r:embed="(rId\d+)"', xml_str)
                        if r_embed_matches:
                            r_id = r_embed_matches[0]
                            media_path = r_id_to_media.get(r_id)
                    except Exception as e:
                        # Fallback: continue without rId
                        pass
                    
                    figures.append(FigureCandidate(
                        para_idx=content_idx,
                        run_idx=run_idx,
                        width_emu=width_emu,
                        height_emu=height_emu,
                        r_id=r_id,
                        media_path=media_path,
                        asset_id=None,  # Will be set after extraction
                        filename=None   # Will be set after extraction
                    ))
                paragraph_counter += 1
            else:
                # Empty paragraph
                para_texts.append("")
            content_idx += 1
                
        elif element.tag.endswith('}tbl'):  # Table
            # Extract table text content
            table_text = ""
            try:
                # Get text from table cells more accurately to avoid duplication
                cell_texts = []
                # Only get text from w:t (text) elements to avoid duplication
                for t_elem in element.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    if t_elem.text and t_elem.text.strip():
                        cell_texts.append(t_elem.text.strip())
                table_text = " ".join(cell_texts).strip()
            except Exception as e:
                # Fallback to simpler method if xpath fails
                try:
                    cell_texts = []
                    for cell in element.iter():
                        if cell.tag.endswith('}t') and cell.text and cell.text.strip():
                            cell_texts.append(cell.text.strip())
                    table_text = " ".join(cell_texts).strip()
                except:
                    table_text = ""
            
            # Convert traditional to simplified if converter provided
            if converter and table_text:
                table_text = safe_traditional_to_simplified(table_text, converter)
            
            # Add table text to para_texts
            para_texts.append(table_text)
            
            # Find drawings in this table
            table_drawings = element.xpath('.//w:drawing')
            for run_idx, drawing in enumerate(table_drawings):
                # Extract dimensions if available
                extents = drawing.xpath('.//wp:extent')
                width_emu = height_emu = None
                if extents:
                    extent = extents[0]
                    width_emu = int(extent.get('cx', 0))
                    height_emu = int(extent.get('cy', 0))
                
                # Extract rId from drawing XML
                r_id = None
                media_path = None
                try:
                    xml_str = drawing.xml
                    r_embed_matches = re.findall(r'r:embed="(rId\d+)"', xml_str)
                    if r_embed_matches:
                        r_id = r_embed_matches[0]
                        media_path = r_id_to_media.get(r_id)
                except Exception as e:
                    # Fallback: continue without rId
                    pass
                
                figures.append(FigureCandidate(
                    para_idx=content_idx,
                    run_idx=run_idx,
                    width_emu=width_emu,
                    height_emu=height_emu,
                    r_id=r_id,
                    media_path=media_path,
                    asset_id=None,  # Will be set after extraction
                    filename=None   # Will be set after extraction
                ))
            content_idx += 1
    
    return figures, para_texts, page_width_emu

# --------- Image Asset Extraction ----------
def extract_and_hash_images(figures: List[FigureCandidate], docx_path: str, assets_dir: str) -> Dict[str, Any]:
    """Extract images from DOCX and calculate real SHA256 hashes"""
    assets = {}
    
    # Ensure assets directory exists
    os.makedirs(assets_dir, exist_ok=True)
    
    # Extract images from DOCX ZIP
    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_file:
            for figure in figures:
                if not figure.media_path:
                    continue
                    
                # Full path in ZIP: word/media/imageN.ext
                zip_media_path = f"word/{figure.media_path}"
                
                if zip_media_path not in zip_file.namelist():
                    continue
                
                # Extract image data
                image_data = zip_file.read(zip_media_path)
                
                # Calculate real SHA256
                sha256_hash = hashlib.sha256(image_data).hexdigest()
                asset_id = f"img_{sha256_hash[:12]}"
                
                # Generate output filename (preserve extension)
                _, ext = os.path.splitext(figure.media_path)
                output_filename = f"{asset_id}{ext}"
                output_path = os.path.join(assets_dir, output_filename)
                
                # Write image to disk
                with open(output_path, 'wb') as f:
                    f.write(image_data)
                
                # Update figure with real asset info
                figure.asset_id = asset_id
                figure.filename = output_path
                
                # Store asset metadata
                if asset_id not in assets:
                    assets[asset_id] = {
                        "asset_id": asset_id,
                        "filename": os.path.join(os.path.basename(assets_dir), output_filename),
                        "sha256": sha256_hash
                    }
    
    except Exception as e:
        # Fallback: generate placeholder assets for figures without real extraction
        for i, figure in enumerate(figures):
            if not figure.asset_id:
                placeholder_id = f"img_placeholder_{i:03d}"
                figure.asset_id = placeholder_id
                if placeholder_id not in assets:
                    assets[placeholder_id] = {
                        "asset_id": placeholder_id,
                        "filename": f"assets/media/placeholder_{i:03d}.png",
                        "sha256": placeholder_id + "0" * (64 - len(placeholder_id))
                    }
    
    return assets

# --------- Phase 2: Group Figures ----------
def group_figures(figures: List[FigureCandidate], para_texts: List[str], 
                 page_width_emu: int, config: Config) -> List[GroupCandidate]:
    """Group figures using two-phase algorithm"""
    if not figures:
        return []
    
    groups = []
    used_figure_indices = set()
    
    # Phase 1: Same-paragraph grouping (row layout)
    figures_by_para = {}
    for i, fig in enumerate(figures):
        para_idx = fig.para_idx
        if para_idx not in figures_by_para:
            figures_by_para[para_idx] = []
        figures_by_para[para_idx].append((i, fig))
    
    for para_idx, para_figs in figures_by_para.items():
        if len(para_figs) >= 2:
            # Multiple figures in same paragraph -> row group
            group_figures = [fig for _, fig in para_figs]
            reason = f"row by same-paragraph(para={para_idx}, {len(group_figures)} images)"
            groups.append(GroupCandidate(
                figures=group_figures,
                layout='row',
                reason=reason
            ))
            used_figure_indices.update(i for i, _ in para_figs)
    
    # Phase 2: Adjacent-paragraph grouping (column layout)
    remaining_figures = [(i, fig) for i, fig in enumerate(figures) if i not in used_figure_indices]
    
    i = 0
    while i < len(remaining_figures):
        current_idx, current_fig = remaining_figures[i]
        group_figures = [current_fig]
        group_indices = [current_idx]
        reason_parts = [f"para={current_fig.para_idx}"]
        
        # Look ahead for adjacent figures
        j = i + 1
        while j < len(remaining_figures):
            next_idx, next_fig = remaining_figures[j]
            current_para = group_figures[-1].para_idx
            next_para = next_fig.para_idx
            
            # Check gap between paragraphs
            gap = next_para - current_para - 1
            if gap > config.max_gap_paras:
                break
            
            # Check if there's substantial text between images
            has_substantial_text = False
            for para_idx in range(current_para + 1, next_para):
                text = para_texts[para_idx]
                if text and len(text) > config.max_title_len and not is_credit_line(text):
                    has_substantial_text = True
                    break
            
            if has_substantial_text:
                break
            
            # Determine if this should be row or column layout
            layout = 'column'  # Default for adjacent paragraphs
            if (current_fig.width_emu and next_fig.width_emu and 
                (current_fig.width_emu + next_fig.width_emu) <= config.page_width_ratio * page_width_emu):
                layout = 'row'  # Could be side-by-side despite being in different paras
            
            group_figures.append(next_fig)
            group_indices.append(next_idx)
            reason_parts.append(f"para={next_fig.para_idx}")
            j += 1
        
        # Determine final layout for the group
        if len(group_figures) > 1:
            layout = 'row' if len(set(fig.para_idx for fig in group_figures)) == 1 else 'column'
            reason = f"{layout} by adjacent-paragraphs({', '.join(reason_parts)}, gap≤{config.max_gap_paras})"
        else:
            layout = 'column'  # Single figure
            reason = f"single figure(para={current_fig.para_idx})"
        
        groups.append(GroupCandidate(
            figures=group_figures,
            layout=layout,
            reason=reason
        ))
        
        used_figure_indices.update(group_indices)
        i = j if len(group_figures) > 1 else i + 1
    
    return groups

# --------- Phase 3: Assign Titles and Credits ----------
def assign_titles_and_credits(groups: List[GroupCandidate], para_texts: List[str], config: Config, doc_full_title: str = None, converter=None):
    """Assign titles and credits to groups"""
    
    for group in groups:
        if not group.figures:
            continue
            
        first_fig = group.figures[0]
        last_fig = group.figures[-1]
        
        title = None
        credit = None
        
        if config.use_explicit_markers:
            # New method: Look for explicit markers
            # Find title (look before first figure for "标题" or "Title")
            for offset in [-2, -1]:  # Check before figure
                check_idx = first_fig.para_idx + offset
                if 0 <= check_idx < len(para_texts):
                    text = para_texts[check_idx]
                    if text and is_title_paragraph(text):
                        title_content = extract_title_content(text)
                        if title_content and title_content != doc_full_title:
                            title = safe_traditional_to_simplified(title_content, converter) if converter else title_content
                            break
            
            # Find credit (look after last figure for "来源" or "Source")
            for offset in [1, 2]:  # Check after figure
                check_idx = last_fig.para_idx + offset
                if 0 <= check_idx < len(para_texts):
                    text = para_texts[check_idx]
                    if text and is_source_paragraph(text):
                        credit_content = extract_source_content(text)
                        if credit_content:
                            credit = safe_traditional_to_simplified(credit_content, converter) if converter else credit_content
                            break
        
        # Fallback to legacy heuristic method if explicit markers not found
        if not title:
            for offset in [-2, -1]:  # Check only before first figure (titles should precede images)
                check_idx = first_fig.para_idx + offset
                if 0 <= check_idx < len(para_texts):
                    text = para_texts[check_idx]
                    # Skip document title and explicit markers (to avoid double processing)
                    if (text and text != doc_full_title and 
                        not is_title_paragraph(text) and not is_source_paragraph(text) and
                        is_short_title(text, config.max_title_len)):
                        title = safe_traditional_to_simplified(text, converter) if converter else text
                        break
        
        if not credit:
            for offset in [1, 2]:  # Check only after last figure (sources should follow images)
                check_idx = last_fig.para_idx + offset
                if 0 <= check_idx < len(para_texts):
                    text = para_texts[check_idx]
                    # Skip explicit markers (to avoid double processing)
                    if text and not is_title_paragraph(text) and not is_source_paragraph(text) and is_potential_credit_line(text):
                        credit = normalize_credit(text)
                        if converter:
                            credit = safe_traditional_to_simplified(credit, converter)
                        break
        
        group.title = title
        group.credit = credit
        
        # Update debug reason with method indication
        method_title = "explicit" if (config.use_explicit_markers and title and 
                                    any(is_title_paragraph(para_texts[first_fig.para_idx + offset]) 
                                       for offset in [-2, -1] if 0 <= first_fig.para_idx + offset < len(para_texts) and 
                                       extract_title_content(para_texts[first_fig.para_idx + offset]) == title)) else "heuristic"
        method_credit = "explicit" if (config.use_explicit_markers and credit and 
                                     any(is_source_paragraph(para_texts[last_fig.para_idx + offset]) 
                                        for offset in [1, 2] if 0 <= last_fig.para_idx + offset < len(para_texts) and 
                                        extract_source_content(para_texts[last_fig.para_idx + offset]) == credit)) else "heuristic"
        
        title_part = f"title({method_title}): '{title[:20]}...'" if title else "title: None"
        credit_part = f"credit({method_credit}): '{credit[:20]}...'" if credit else "credit: None"
        group.reason += f", {title_part}, {credit_part}"

# --------- Main Conversion ----------
def convert_docx_to_ncj(docx_path: str, config: Config) -> Dict[str, Any]:
    """Convert DOCX to NCJ format with improved figure grouping"""
    doc = Document(docx_path)
    
    # Initialize traditional to simplified converter if needed
    converter = None
    if config.traditional_to_simplified:
        try:
            converter = opencc.OpenCC('t2s')  # Traditional to Simplified
        except Exception as e:
            print(f"Warning: Failed to initialize OpenCC converter: {e}", file=sys.stderr)
    
    # Extract figures and paragraph texts
    figures, para_texts, page_width_emu = extract_figures_from_docx(doc, docx_path, converter)
    
    # Extract and hash real image assets
    assets = extract_and_hash_images(figures, docx_path, config.assets_dir)
    
    # Group figures
    groups = group_figures(figures, para_texts, page_width_emu, config)
    
    # Extract document metadata
    doc_title = None
    doc_date = None
    doc_full_title = None  # Store full title for skipping logic
    if para_texts:
        first_text = para_texts[0]
        match = DOC_TITLE_RE.match(first_text or '')
        if match:
            doc_full_title = first_text.strip()  # Full title for comparison
            doc_title = match.group(2).strip()   # Clean title without date
            if converter:
                doc_title = safe_traditional_to_simplified(doc_title, converter)
            doc_date = parse_date_from_yyMMdd(match.group(1))
    
    # Assign titles and credits
    assign_titles_and_credits(groups, para_texts, config, doc_full_title, converter)
    
    # Build output blocks
    out_blocks = []
    debug_info = []
    
    # Add non-image paragraphs and figures
    para_consumed = set()
    
    # First pass: mark paragraphs that contain figures or are consumed as titles/credits
    for group in groups:
        for fig in group.figures:
            para_consumed.add(fig.para_idx)
        
        # Mark title paragraph as consumed (check both original and extracted content)
        if group.title:
            for para_idx, text in enumerate(para_texts):
                # Check if this paragraph is a title paragraph with matching content
                if (is_title_paragraph(text) and extract_title_content(text) == group.title) or text == group.title:
                    para_consumed.add(para_idx)
                    break
        
        # Mark credit paragraph as consumed (check both original and extracted content)
        if group.credit:
            # Find the closest credit paragraph to the last figure in this group
            last_fig = group.figures[-1]
            best_match_idx = None
            best_distance = float('inf')
            
            for para_idx, text in enumerate(para_texts):
                # Check if this paragraph is a source paragraph with matching content
                # or if normalizing this text produces the same credit
                if (is_source_paragraph(text) and extract_source_content(text) == group.credit) or \
                   (text and normalize_credit(text) == group.credit):
                    distance = abs(para_idx - last_fig.para_idx)
                    if distance < best_distance:
                        best_distance = distance
                        best_match_idx = para_idx
            
            # Mark the closest match as consumed
            if best_match_idx is not None:
                para_consumed.add(best_match_idx)
                if config.debug:
                    debug_info.append(f"Consumed closest para {best_match_idx}: '{para_texts[best_match_idx][:30]}...' -> credit: '{group.credit}' (distance: {best_distance})")
    
    # Process all content in document order
    group_counter = 1
    figure_global_index = 0
    processed_groups = set()
    
    for para_idx, para_text in enumerate(para_texts):
        # Skip document title
        if para_idx == 0 and doc_full_title and para_text == doc_full_title:
            continue
            
        # Check if this paragraph starts a new group
        group_starting_here = None
        for g in groups:
            if g.figures and g.figures[0].para_idx == para_idx and id(g) not in processed_groups:
                group_starting_here = g
                break
        
        if group_starting_here:
            # Output the entire group
            group_id = f"grp_{group_counter:04d}"
            group_len = len(group_starting_here.figures)
            
            # Generate figures for this group
            for seq, fig in enumerate(group_starting_here.figures):
                # Use real asset_id from extracted images
                asset_id = fig.asset_id or f"img_missing_{figure_global_index:03d}"
                
                # Create figure block
                figure_block = {
                    "type": "figure",
                    "image": {"asset_id": asset_id},
                    "title": group_starting_here.title if seq == 0 else None,  # Title on first figure
                    "credit": group_starting_here.credit if seq == group_len - 1 else None,  # Credit on last figure
                    "group_id": group_id,
                    "group_seq": seq + 1,
                    "group_len": group_len,
                    "layout": group_starting_here.layout
                }
                
                out_blocks.append(figure_block)
                figure_global_index += 1
            
            debug_info.append(f"{group_id}: {group_starting_here.reason}")
            processed_groups.add(id(group_starting_here))
            group_counter += 1
            
        elif para_text and para_idx not in para_consumed:
            # Regular text paragraph
            # Convert text if needed (should already be converted, but double-check)
            display_text = para_text
            out_blocks.append({
                "type": "paragraph", 
                "text": display_text
            })
    
    # Build final NCJ structure
    ncj = {
        "doc": {
            "title": doc_title,
            "date": doc_date,
            "locale": "zh-CN",
            "version": "v1",
            "source_file": os.path.basename(docx_path)
        },
        "blocks": out_blocks,
        "assets": list(assets.values()),
        "report": {
            "warnings": [],
            "debug": debug_info if config.debug else []
        }
    }
    
    return ncj

def main():
    import argparse
    
    parser = argparse.ArgumentParser(description='Convert DOCX to NCJ with advanced figure grouping')
    parser.add_argument('input', help='Input DOCX file')
    parser.add_argument('--out', default='content.json', help='Output JSON file')
    parser.add_argument('--assets-dir', default='assets/media', 
                       help='Directory to extract image assets')
    parser.add_argument('--max_title_len', type=int, default=60, 
                       help='Maximum length for title detection')
    parser.add_argument('--max_gap_paras', type=int, default=1,
                       help='Maximum paragraphs gap for grouping')
    parser.add_argument('--page_width_ratio', type=float, default=0.95,
                       help='Page width ratio for row layout detection')
    parser.add_argument('--debug', action='store_true',
                       help='Include debug information in output')
    parser.add_argument('--traditional-to-simplified', action='store_true',
                       help='Convert traditional Chinese to simplified Chinese')
    parser.add_argument('--no-explicit-markers', action='store_true',
                       help='Disable explicit title/source marker detection (use legacy heuristic only)')
    
    args = parser.parse_args()
    
    # Configure
    config = Config()
    config.max_title_len = args.max_title_len
    config.max_gap_paras = args.max_gap_paras
    config.page_width_ratio = args.page_width_ratio
    config.debug = args.debug
    config.assets_dir = args.assets_dir
    config.traditional_to_simplified = args.traditional_to_simplified
    config.use_explicit_markers = not args.no_explicit_markers  # Default: use explicit markers
    
    # Convert
    try:
        ncj = convert_docx_to_ncj(args.input, config)
        
        # Output
        if args.out == '-':
            json.dump(ncj, sys.stdout, ensure_ascii=False, indent=2)
        else:
            with open(args.out, 'w', encoding='utf-8') as f:
                json.dump(ncj, f, ensure_ascii=False, indent=2)
                
        print(f"Converted {args.input} -> {args.out}")
        
        # Print summary
        total_figures = len([b for b in ncj['blocks'] if b['type'] == 'figure'])
        groups = len(set(b.get('group_id') for b in ncj['blocks'] if b['type'] == 'figure' and b.get('group_id')))
        multi_groups = len([g for g in set(b.get('group_id') for b in ncj['blocks'] if b['type'] == 'figure' and b.get('group_id')) 
                           if any(b.get('group_len', 1) > 1 for b in ncj['blocks'] if b.get('group_id') == g)])
        
        print(f"Summary: {total_figures} figures, {groups} groups, {multi_groups} multi-figure groups")
        
        if config.debug:
            print("Debug info:")
            for info in ncj['report']['debug']:
                print(f"  {info}")
    
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    main()